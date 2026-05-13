#!/usr/bin/env python3
"""
create_sample_cv.py – Generate a sample academic CV .docx for testing.

Usage:
    python3 scripts/cv/create_sample_cv.py

Writes public/cv.docx. Replace it with your real CV when ready.
"""
import sys
from pathlib import Path


def require_docx():
    try:
        import docx
        return docx
    except ImportError:
        print(
            "ERROR: python-docx is not installed.\n"
            "Install it with:  pip install python-docx",
            file=sys.stderr,
        )
        sys.exit(1)


def main():
    docx = require_docx()
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Name / contact block
    name_para = doc.add_paragraph()
    run = name_para.add_run("Dr. Alex Sample")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph("Assistant Professor of Economics")
    doc.add_paragraph("Department of Economics, Sample University")
    doc.add_paragraph("email: alex.sample@sample.edu  |  https://alexsample.example.edu")

    doc.add_paragraph("")  # spacer

    # Research Areas
    doc.add_heading("Research Areas", level=1)
    doc.add_paragraph("Political Economy; Comparative Politics; Public Finance; Electoral Systems")

    # Employment
    doc.add_heading("Employment", level=1)
    doc.add_paragraph("Assistant Professor, Sample University, 2021–present")
    doc.add_paragraph("Postdoctoral Fellow, Institute for Advanced Study, 2019–2021")

    # Education
    doc.add_heading("Education", level=1)
    doc.add_paragraph("Ph.D., Political Science, State University, 2019")
    doc.add_paragraph("M.A., Political Science, State University, 2015")
    doc.add_paragraph("B.A., Economics and Government, Liberal Arts College, 2013")

    # Publications
    doc.add_heading("Publications", level=1)
    doc.add_paragraph(
        'Sample, Alex. "Fiscal Policy and Electoral Accountability." '
        'Journal of Politics 85(2): 410–425. 2023.'
    )
    doc.add_paragraph(
        'Sample, Alex, and Jordan Co-Author. "Redistribution under Federalism." '
        'American Political Science Review 117(1): 55–72. 2022.'
    )
    doc.add_paragraph(
        'Sample, Alex. "Budget Transparency and Voter Learning." '
        'Quarterly Journal of Political Science 16(3): 301–330. 2021.'
    )

    # Work in Progress
    doc.add_heading("Work in Progress", level=1)
    doc.add_paragraph(
        '"Tax Salience and Policy Preferences." '
        'Under review at the Journal of Public Economics.'
    )
    doc.add_paragraph(
        '"Intergovernmental Transfers and Local Accountability." '
        'Manuscript in preparation.'
    )
    doc.add_paragraph(
        '"Electoral Rules and Distributive Politics." '
        'Data collection in progress.'
    )

    # Teaching
    doc.add_heading("Teaching", level=1)
    doc.add_paragraph("Introduction to Political Economy (undergraduate), Sample University, 2022, 2023")
    doc.add_paragraph("Public Finance and Public Policy (graduate), Sample University, 2022")
    doc.add_paragraph("Comparative Political Economy (undergraduate), Sample University, 2021")

    # Invited Talks
    doc.add_heading("Invited Talks", level=1)
    doc.add_paragraph('"Fiscal Policy and Elections." Annual Conference on Political Economy, 2024.')
    doc.add_paragraph('"Tax Design and Voter Behavior." Workshop on Public Finance, 2023.')

    # Honors and Grants
    doc.add_heading("Honors and Awards", level=1)
    doc.add_paragraph("National Science Foundation Grant, #SES-0000001, 2022–2024")
    doc.add_paragraph("Best Paper Award, Comparative Politics Section, APSA, 2022")
    doc.add_paragraph("Graduate Research Fellowship, State University, 2016–2019")

    # Service
    doc.add_heading("Service", level=1)
    doc.add_paragraph("Referee: Journal of Politics, APSR, QJPS, CPS")
    doc.add_paragraph("Organiser, Graduate Workshop in Political Economy, Sample University, 2022–present")

    output = Path("public/cv.docx")
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    print(f"Sample CV written to {output}")
    print("Replace it with your real CV when ready.")


if __name__ == "__main__":
    main()
